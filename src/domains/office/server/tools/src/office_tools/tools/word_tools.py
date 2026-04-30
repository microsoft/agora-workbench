"""
Word document extraction tools.

All functions receive a ``file`` parameter that is a ``Path`` to a cached
Word file on disk (resolved by AssetResolutionMiddleware).
"""

from pathlib import Path
from typing import Any, Union

import olefile
from docx import Document

_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def _check_irm(file: Path) -> None:
    """Raise ``ValueError`` if *file* is an OLE2 container (IRM-encrypted or legacy .doc)."""
    with open(file, "rb") as f:
        header = f.read(8)
    if header != _OLE2_MAGIC:
        return
    if olefile.isOleFile(str(file)):
        with olefile.OleFileIO(str(file)) as ole:
            if ole.exists("\x09DRMContent"):
                raise ValueError(
                    f"File {file.name} is still IRM-encrypted. "
                    "IRM decryption may not have run — check that "
                    "the server has valid Azure RMS credentials."
                )
    raise ValueError(
        f"File {file.name} is a legacy OLE2 document (.doc format). "
        "Only .docx files are supported. Please re-save the file as .docx."
    )


def extract_word_text(file: Union[Path, Any]) -> dict:
    """
    Extract all text content from a Word document.

    Args:
        file: Path to the .docx file.

    Returns:
        Dict with ``text`` (paragraphs joined by newlines) and ``paragraph_count``.
    """
    file = Path(file)
    _check_irm(file)
    doc = Document(str(file))
    paragraphs = [p.text for p in doc.paragraphs]
    return {
        "text": "\n".join(paragraphs),
        "paragraph_count": len(paragraphs),
    }


def extract_word_tables(file: Union[Path, Any]) -> dict:
    """
    Extract all tables from a Word document.

    Each table uses its first row as column headers. Subsequent rows are
    returned as dicts keyed by those headers.

    Args:
        file: Path to the .docx file.

    Returns:
        Dict with ``tables`` (list of list-of-row-dicts) and ``table_count``.
    """
    file = Path(file)
    _check_irm(file)
    doc = Document(str(file))

    tables = []
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        headers = rows[0] if rows else []
        data_rows = [dict(zip(headers, row)) for row in rows[1:]]
        tables.append({"headers": headers, "rows": data_rows})

    return {
        "tables": tables,
        "table_count": len(tables),
    }
